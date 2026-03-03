"""
Word (.docx) document parser using python-docx.
Extracts paragraphs, tables, and metadata into a structured format
suitable for AI-assisted financial analysis.
"""
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document
from docx.table import Table as DocxTable

logger = logging.getLogger(__name__)


@dataclass
class DocxTableData:
    """A single table extracted from a Word document."""
    headers: list[str]
    rows: list[list[str]]

    @property
    def row_count(self) -> int:
        return len(self.rows)

    def to_text(self) -> str:
        lines = []
        if self.headers:
            lines.append(" | ".join(self.headers))
            lines.append("-" * len(lines[0]))
        for row in self.rows:
            lines.append(" | ".join(row))
        return "\n".join(lines)


@dataclass
class DocxDocument:
    """Structured representation of a parsed Word document."""
    filepath: str
    paragraphs: list[str]
    tables: list[DocxTableData]
    metadata: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        """All paragraph text joined with newlines."""
        return "\n".join(p for p in self.paragraphs if p.strip())

    @property
    def full_text_with_tables(self) -> str:
        """Paragraph text plus table data rendered as text blocks."""
        parts = [self.full_text]
        for i, table in enumerate(self.tables):
            parts.append(f"\n--- TABLE {i + 1} ---\n{table.to_text()}")
        return "\n".join(parts)

    @property
    def has_tables(self) -> bool:
        return len(self.tables) > 0


def _extract_table(table: DocxTable) -> Optional[DocxTableData]:
    """Extract a python-docx Table into a DocxTableData."""
    rows_raw: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows_raw.append(cells)

    if not rows_raw:
        return None

    # Treat first row as headers if it looks like a header row
    # (all non-empty, no numeric-only cells)
    first_row = rows_raw[0]
    is_header = (
        all(cell for cell in first_row)
        and not all(cell.replace(",", "").replace(".", "").replace("$", "").isdigit()
                     for cell in first_row if cell)
    )

    if is_header and len(rows_raw) > 1:
        return DocxTableData(headers=first_row, rows=rows_raw[1:])

    return DocxTableData(headers=[], rows=rows_raw)


def extract_docx(filepath: str) -> DocxDocument:
    """
    Parse a .docx file and return a DocxDocument with paragraphs, tables,
    and core metadata.

    Raises FileNotFoundError if the file does not exist.
    """
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {filepath}")

    doc = Document(filepath)

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    tables: list[DocxTableData] = []
    for table in doc.tables:
        try:
            extracted = _extract_table(table)
            if extracted and (extracted.headers or extracted.rows):
                tables.append(extracted)
        except Exception as e:
            logger.warning(f"Table extraction failed in {path.name}: {e}")

    metadata: dict = {}
    core = doc.core_properties
    if core.title:
        metadata["title"] = core.title
    if core.author:
        metadata["author"] = core.author
    if core.created:
        metadata["created"] = core.created.isoformat()
    if core.modified:
        metadata["modified"] = core.modified.isoformat()

    logger.info(
        f"Extracted {len(paragraphs)} paragraphs, {len(tables)} tables from {path.name}"
    )
    return DocxDocument(
        filepath=filepath,
        paragraphs=paragraphs,
        tables=tables,
        metadata=metadata,
    )
